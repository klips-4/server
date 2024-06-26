import ast
import json
from Model.Client import Client
from sqlalchemy import desc, text
from sqlalchemy.orm import Session
from flask import jsonify
from datetime import datetime, timedelta

from Helpers.httpResponse import HttpResponse
from Model import BaseModel
from app import engine, db
from serializator import *

import docx
import pandas as pd

from sqlalchemy_paginator import Paginator


class BaseClass:
    USE_NAVIGATION: bool = True
    FIELD_SORT: str = None
    AREA: str = 'BaseClass'

    _session: Session = engine.session
    _additional_methods: dict = None

    def __init__(self):
        self._methods_maps = {
            'Create': self.create,
            'Get': self.get,
            'Delete': self.delete,
            'Update': self.update,
            'List': self.list
        }

        if self._additional_methods:
            self.methods_map.update(self._additional_methods)

    @property
    def methods_map(self):
        """
        Получение доступных методов сущности
        :return: Список методов сущностей
        """
        return self._methods_maps

    @staticmethod
    def get_model(new_model: bool = False) -> BaseModel:
        """
        Получение модели

        :param new_model: Признак создавать экземпляр или нет
        :return: Модель сущности
        """
        return BaseModel() if new_model else BaseModel

    def get(self, **kwargs):
        """
        Получение записи по ключу
        :param kwargs: В date содержится ключ записи
        :return: запись или ошибка
        """
        key = kwargs.get('data')
        query = self._session.query(self.get_model().get(key))

        if query:
            return HttpResponse.make(data=query.to_dict())
        else:
            return HttpResponse.make(success=False, error_text=f"Запись в таблице {self.get_model(True).__tablename__}"
                                                               f"по ключу {key} не найдена")

    def list(self, **kwargs):
        """
        Метод получения списка с фильтрацией

        :param self: Параметры с фильтром
        :param kwargs:
        :return: список записей

        """
        data = kwargs.get('filter')
        number = data.get('0')

        if number == None:
            number = 1

        lastEmployeeID = int(number)

        if lastEmployeeID == 1:
            list = '''
                SELECT * FROM public.clients
                WHERE ID > 0
                ORDER BY id ASC
                LIMIT 5;
                '''

        if lastEmployeeID == 2:
            list = '''
                       SELECT * FROM public.clients
                       WHERE ID > 5 AND ID < 11
                       ORDER BY id ASC
                       LIMIT 5;
                       '''

        if lastEmployeeID == 3:
            list = '''
                              SELECT * FROM public.clients
                              WHERE ID > 11
                              ORDER BY id ASC
                              LIMIT 5;
                              '''

        result = db.session.execute(text(list))

        self.clientsData = kwargs.get('data') or []
        entities = db.session.execute(text(list))

        return HttpResponse.make(data=to_dict(entities))

    def create(self, **kwargs):
        """
        Метод получения формата записи
        :param kwargs: запись с заполнеными полями
        :return: формат записи с предустановленными полями
        """
        record = kwargs.get('data')
        format = "YYYY-mm-dd HH:MM:SS"
        date1 = datetime.strptime(record[0], '%Y-%m-%dT%H:%M:%S.%fZ')
        date2 = datetime.strptime(record[1], '%Y-%m-%dT%H:%M:%S.%fZ')
        querys = '''
        WITH done_counts AS (
          SELECT nature_of_appeal, COUNT(*) AS acount
          FROM public.clients
          WHERE client_date_application BETWEEN '%s' AND '%s'
            AND public.clients.result = 'Выявлено'
          GROUP BY nature_of_appeal
        ),
        process_counts AS (
          SELECT nature_of_appeal, COUNT(*) AS bcount
          FROM public.clients
          WHERE client_date_application BETWEEN '%s' AND '%s'
            AND public.clients.result = 'Не выявлено'
          GROUP BY nature_of_appeal
        )
        SELECT
          COALESCE(a.nature_of_appeal, b.nature_of_appeal) AS nature_of_appeal,
          a.acount,
          b.bcount
        FROM done_counts a
        FULL OUTER JOIN process_counts b ON a.nature_of_appeal = b.nature_of_appeal
         ''' % (date1, date2, date1, date2)

        result = db.session.execute(text(querys))

        results_list = [dict(zip(result.keys(), row)) for row in result]

        doc = docx.Document()
        doc.add_heading('Отчет', 0)

        table = doc.add_table(rows=1, cols=3)
        row = table.rows[0].cells
        row[0].text = ''
        row[1].text = 'Выявлено'
        row[2].text = 'Не выявлено'

        for i in results_list:
            row = table.add_row().cells
            row[0].text = 'Характер выявленных нарушений' + ' ' + str(i.get('nature_of_appeal'))
            if str(i.get('acount')) == 'None':
                row[1].text = '-'
            else:
                row[1].text = str(i.get('acount'))
            if str(i.get('bcount')) == 'None':
                row[2].text = '-'
            else:
                row[2].text = str(i.get('bcount'))
        doc.save('report.docx')

        return 'HttpResponse.make(data=data)'

    def update(self, **kwargs):

        """
        Обновление записи или создание записи
        :param kwargs: запись для сохранения
        :return: Обновленная запись
        """

        self.clientsData = kwargs.get('data') or []
        self.active = self.clientsData.get('active')
        self.clientsId = self.clientsData.get('row')
        self.color = self.clientsData.get('color')
        lst = ast.literal_eval(self.clientsId)

        for client_id in lst:
            client = Client.query.get(client_id)
            model = self._session.query(self.get_model()).get(client_id)

            if client:
                client.result = self.active
                client.color = self.color
                db.session.commit()

        entities = Client.query.order_by(Client.id).all()

        return HttpResponse.make(data=to_dict(entities))

        # record = kwargs.get('data')
        #
        # if isinstance(record, str):
        #     record = json.loads(record)
        #
        # if not record:
        #     return HttpResponse.make(success=False, error_text="Не переданные данные")
        #
        # if not record.get('id'):
        #     return self._new(record)
        # else:
        #     return self._update(record)

    def delete(self, **kwargs):
        """
        Удаление записи
        :param kwargs: ключ записи
        :return: результат удаления
        """
        key = kwargs.get('data')
        query = self._session.query(self.get_model()).get(key)

        if query:
            self._session.delete(query)
            self._session.commit()

            return HttpResponse.make()

        else:
            return HttpResponse.make(success=False, error_text="Не найдена запись по ключу")

    def _prepare_list(self, result, filter_params):
        """
        Постобработка перед выдачей результата

        :param result: Список сущностей
        :param filter_params: фильтр
        :return: список обработанных записей
        """
        return result

    def _prepare_list_result(self, filter_params):
        """
        Метод получения результата с фильтрацией

        :param filter_params: Фильтр
        :return: Список записей
        """
        result = []

        query = self._prepare_query_filter(self._session.query(self.get_model()), filter_params)

        query.order_by(Client.client_name)

        count_result = query.count() if query else 0

        if count_result:
            result = [item.to_dict() for item in query.all()]

        return result

    def _prepare_query_filter(self, query, filter_params):
        """
         Применение фильтрации к объекту записи
        :param query: Объект запроса
        :param filter_params: Фильтр
        :return: Объект запроса с фильтрацией
        """
        return query

    def _new(self, record):
        """
        Создание новой записи в БД
        :param record: Данные которые передаем в модель
        :return: Обновленная запись в БД
        """
        model = self.get_model(True)
        model.from_object(record)
        self._session.add_all([model])
        self._session.commit()

        return HttpResponse.make(data=model.to_dict())

    def _update(self, record):
        """
        Обновление записи
        :param record: Обновляемые данные для модели
        :return: Обновленная запись в БД
        """
        model = self._session.query(self.get_model()).get(record.get('id'))
        print(model)

        if model:
            model.from_object(record)
            self._session.add_all([model])
            self._session.commit()

            return HttpResponse.make(data=model.to_dict())
        else:
            return HttpResponse.make(success=False, error_text="Нет записи для обновления БД")
